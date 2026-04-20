from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

import openpyxl
import xlrd
from docx import Document
from pypdf import PdfReader


class ExtractionError(RuntimeError):
    """Raised when a file cannot be converted into text."""


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".doc":
        return _extract_doc(path)
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    if suffix == ".xls":
        return _extract_xls(path)
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")

    raise ExtractionError(f"Unsupported file type: {suffix or 'unknown'}")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _clean_text("\n".join(parts))


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))

    return _clean_text("\n".join(lines))


def _extract_doc(path: Path) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        raise ExtractionError("Legacy .doc extraction requires antiword on PATH.")

    result = subprocess.run(
        [antiword, str(path)],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        message = result.stderr.strip() or result.stdout.strip() or "antiword failed"
        raise ExtractionError(message)
    return _clean_text(result.stdout)


def _extract_xlsx(path: Path) -> str:
    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = []

    for sheet in workbook.worksheets:
        lines.append(f"[Sheet] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                lines.append(" | ".join(values))

    return _clean_text("\n".join(lines))


def _extract_xls(path: Path) -> str:
    workbook = xlrd.open_workbook(path)
    lines: list[str] = []

    for sheet in workbook.sheets():
        lines.append(f"[Sheet] {sheet.name}")
        for row_index in range(sheet.nrows):
            values = [str(value).strip() for value in sheet.row_values(row_index) if str(value).strip()]
            if values:
                lines.append(" | ".join(values))

    return _clean_text("\n".join(lines))


def _clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
